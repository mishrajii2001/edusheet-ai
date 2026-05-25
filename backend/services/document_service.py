from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import uuid

ALIGNMENT_MAP = {
    "justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT
}

def add_heading(doc, text, font_name, font_size, alignment):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = True
    paragraph.alignment = ALIGNMENT_MAP.get(alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
    return paragraph

def add_body(doc, text, font_name, font_size, alignment):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = False
    paragraph.alignment = ALIGNMENT_MAP.get(alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
    return paragraph

def add_code_block(doc, code_text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(11)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F0F0F0")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:val"), "clear")
    pPr.append(shading)
    return paragraph

def add_image_to_doc(doc, image_path, caption=None, max_width_inches=5.5):
    try:
        from PIL import Image as PILImage
        img = PILImage.open(image_path)
        img_width, img_height = img.size
        aspect = img_width / img_height
        if aspect > 2:
            width = Inches(max_width_inches)
        elif aspect < 0.7:
            width = Inches(max_width_inches * 0.45)
        else:
            width = Inches(max_width_inches * 0.75)
        doc.add_picture(image_path, width=width)
        img_para = doc.paragraphs[-1]
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.size = Pt(10)
                run.font.italic = True
        doc.add_paragraph()
    except Exception as e:
        doc.add_paragraph(f"[Image could not be added: {str(e)}]")

def create_worksheet(content, formatting, images=None, header_image=None):
    doc = Document()

    font_name = formatting.get("font_family", "Times New Roman")
    heading_size = formatting.get("heading_size", 14)
    body_size = formatting.get("body_size", 12)
    alignment = formatting.get("alignment", "justified")
    margin = formatting.get("margin", 1.0)

    # Set margins FIRST
    for section in doc.sections:
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)

    # THEN add header image
    if header_image:
        try:
            doc.add_picture(header_image, width=Inches(6.5))
            header_para = doc.paragraphs[-1]
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        except Exception as e:
            print(f"Header image failed: {str(e)}")

    # THEN add title
    if content.get("title"):
        title = doc.add_paragraph()
        run = title.add_run(content["title"].upper())
        run.font.name = font_name
        run.font.size = Pt(heading_size + 2)
        run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # ... rest of the code stays same

    predefined_sections = [
    ("aim", "Aim"),
    ("introduction", "Introduction"),
    ("objective", "Objective"),
    ("apparatus", "Apparatus"),
    ("task_to_be_done", "Task to be Done"),
    ("theory", "Theory"),
    ("algorithm", "Algorithm"),
    ("code", "Code"),
    ("expected_output", "Expected Output"),
    ("result", "Result"),
    ("conclusion", "Conclusion"),
    ("learning_outcomes", "Learning Outcomes"),
    ("viva_questions", "Viva Questions"),
    ("references", "References")
    ]

    predefined_keys = [k for k, _ in predefined_sections]

    extra_sections = [
        (k, k.replace("_", " ").title())
        for k in content.keys()
        if k not in predefined_keys and k != "title" and content.get(k)
    ]

    sections_order = predefined_sections + extra_sections

    for key, label in sections_order:
        value = content.get(key)
        if not value:
            continue

        add_heading(doc, label, font_name, heading_size, alignment)

        if key in ["code", "expected_output"]:
            clean_value = value.replace("\\n", "\n")
            add_code_block(doc, clean_value)
        elif isinstance(value, list):
            for i, item in enumerate(value, 1):
                add_body(doc, f"{i}. {item}", font_name, body_size, alignment)
        else:
            add_body(doc, value, font_name, body_size, alignment)

        if images and key in images:
            for img in images[key]:
                add_image_to_doc(doc, img["path"], img.get("caption"))

        doc.add_paragraph()

    output_dir = os.path.join(os.path.dirname(__file__), "..", "generated")
    os.makedirs(output_dir, exist_ok=True)
    file_name = f"worksheet_{uuid.uuid4().hex[:8]}.docx"
    file_path = os.path.join(output_dir, file_name)
    doc.save(file_path)

    return file_name